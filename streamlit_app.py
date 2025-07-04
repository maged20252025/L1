import streamlit as st

import streamlit.components.v1 as components

from docx import Document

from docx.shared import Inches

import re

import uuid

import os

import time

import html

import csv

from io import BytesIO



# ----------------------------------------------------

# إعدادات الصفحة الأساسية

# ----------------------------------------------------

st.set_page_config(

    page_title="القوانين اليمنية بآخر تعديلاتها حتى عام 2025م",

    layout="wide",

    initial_sidebar_state="expanded"

)



# ----------------------------------------------------

# ثوابت ومتغيرات عامة

# ----------------------------------------------------

TRIAL_DURATION = 180  # 3 دقائق بالثواني

TRIAL_USERS_FILE = "trial_users.txt"

DEVICE_ID_FILE = "device_id.txt"

ACTIVATED_FILE = "activated.txt"

ACTIVATION_CODES_FILE = "activation_codes.txt"

LAWS_DIR = "laws"



# ----------------------------------------------------

# دوال المساعدة

# ----------------------------------------------------

def get_device_id():

    if os.path.exists(DEVICE_ID_FILE):

        with open(DEVICE_ID_FILE, "r") as f:

            return f.read().strip()

    new_id = str(uuid.uuid4())

    with open(DEVICE_ID_FILE, "w") as f:

        f.write(new_id)

    return new_id



def get_trial_start(device_id):

    if not os.path.exists(TRIAL_USERS_FILE):

        return None

    with open(TRIAL_USERS_FILE, "r") as f:

        reader = csv.reader(f)

        for row in reader:

            if row and row[0] == device_id:

                return float(row[1])

    return None



def register_trial(device_id):

    if not os.path.exists(TRIAL_USERS_FILE):

        with open(TRIAL_USERS_FILE, "w", newline='') as f:

            pass

    with open(TRIAL_USERS_FILE, "a", newline='') as f:

        writer = csv.writer(f)

        writer.writerow([device_id, time.time()])



def is_activated():

    return os.path.exists(ACTIVATED_FILE)



def activate_app(code):

    if not os.path.exists(ACTIVATION_CODES_FILE):

        return False

    with open(ACTIVATION_CODES_FILE, "r") as f:

        codes = [line.strip() for line in f.readlines()]

    if code in codes:

        codes.remove(code)

        with open(ACTIVATION_CODES_FILE, "w") as f:

            for c in codes:

                f.write(c + "\n")

        with open(ACTIVATED_FILE, "w") as f:

            f.write("activated")

        return True

    return False



def highlight_keywords(text, keywords):

    for kw in keywords:

        text = re.sub(f"({re.escape(kw)})", r"<mark>\1</mark>", text, flags=re.IGNORECASE)

    return text



def export_results_to_word(results, filename="نتائج_البحث.docx"):

    document = Document()

    document.add_heading('نتائج البحث في القوانين اليمنية', level=1)

    if not results:

        document.add_paragraph("لم يتم العثور على نتائج للكلمات المفتاحية المحددة.")

    else:

        for i, r in enumerate(results):

            document.add_heading(f"القانون: {r['law']} - المادة: {r['num']}", level=2)

            document.add_paragraph(r['plain'])

            if i < len(results) - 1:

                document.add_page_break()

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()



def normalize_arabic_numbers(text):

    # تحويل الأرقام العربية إلى إنجليزية

    arabic_to_english = str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789')

    return text.translate(arabic_to_english)



# ----------------------------------------------------

# وظيفة التطبيق الرئيسية (بعد التفعيل أو بدء التجربة)

# ----------------------------------------------------

def run_main_app():

    # إضافة CSS لتصحيح اتجاه مربع النص وزر التصدير والعداد

    components.html("""

    <style>

    .scroll-btn {

        position: fixed;

        left: 10px;

        padding: 12px;

        font-size: 24px;

        border-radius: 50%;

        background-color: #c5e1a5;

        color: black;

        cursor: pointer;

        z-index: 9999;

        border: none;

        box-shadow: 1px 1px 5px #888;

    }

    #scroll-top-btn { bottom: 80px; }

    #scroll-bottom-btn { bottom: 20px; }

    /* ---- تخصيص المحاذاة لليمين للـ Metric والـ Download button ---- */

    .rtl-metric {

        direction: rtl;

        text-align: right !important;

        margin-right: 0 !important;

    }

    .rtl-metric .stMetric {

        text-align: right !important;

        direction: rtl;

    }

    .rtl-metric .stMetricDelta {

        display: block !important;

        text-align: right !important;

        direction: rtl;

    }

    .rtl-download-btn {

        direction: rtl;

        text-align: right !important;

        margin-right: 0 !important;

        display: flex;

        flex-direction: row-reverse;

        justify-content: flex-start;

    }

    /* --------- اجبار مربعات النصوص للكتابة من اليمين -------- */

    textarea, .stTextArea textarea {

        direction: rtl !important;

        text-align: right !important;

    }

    /* --------- اجبار كل عناصر النتائج أن تكون يمين -------- */

    .stButton, .stDownloadButton, .stMetric {

        direction: rtl !important;

        text-align: right !important;

    }

    </style>

    <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>⬆️</button>

    <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>⬇️</button>

    """, height=1)



    if not os.path.exists(LAWS_DIR):

        st.error(f"⚠️ مجلد '{LAWS_DIR}/' غير موجود. يرجى التأكد من وجود ملفات القوانين.")

        return



    files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]

    if not files:

        st.warning(f"📂 لا توجد ملفات قوانين في مجلد '{LAWS_DIR}/'.")

        return



    # -- نموذج البحث بمحاذاة يمين --

    st.markdown("""

        <div style="direction: rtl; text-align: right;">

        <h3 style="display: flex; align-items: center; gap: 10px;">🔎 نموذج البحث</h3>

        </div>

    """, unsafe_allow_html=True)

    with st.form("main_search_form"):

        # تخصيص تسمية الحقول مع اتجاه يمين

        st.markdown('<div style="direction: rtl; text-align: right;">اختر قانونًا للبحث:</div>', unsafe_allow_html=True)

        selected_file_form = st.selectbox("", ["الكل"] + files, key="main_file_select", label_visibility="collapsed")

        st.markdown('<div style="direction: rtl; text-align: right;">📌 الكلمات المفتاحية (افصل بفاصلة):</div>', unsafe_allow_html=True)

        keywords_form = st.text_area(

            "",

            key="main_keywords_input",

            help="أدخل الكلمات التي تريد البحث عنها، وافصل بينها بفاصلة إذا كانت أكثر من كلمة."

        )

        # مربع رقم المادة

        st.markdown('<div style="direction: rtl; text-align: right;">🔢 رقم المادة (اختياري):</div>', unsafe_allow_html=True)

        article_number_input = st.text_input(

            "",

            key="article_number_input",

            help="أدخل رقم المادة للبحث عنها مباشرة (يمكن استخدام أرقام عربية أو إنجليزية)."

        )

        # زر البحث مع أيقونة يمين

        search_btn_col = st.columns([1, 2, 12])

        with search_btn_col[2]:

            submitted = st.form_submit_button("🔍 بدء البحث", use_container_width=True)



    if "results" not in st.session_state:

        st.session_state.results = []

    if "search_done" not in st.session_state:

        st.session_state.search_done = False



    # تنفيذ البحث فقط إذا تم إرسال النموذج

    if submitted:

        results = []

        search_files = files if selected_file_form == "الكل" else [selected_file_form]

        kw_list = [k.strip() for k in keywords_form.split(",") if k.strip()] if keywords_form else []

        search_by_article = bool(article_number_input.strip())



        norm_article = normalize_arabic_numbers(article_number_input.strip()) if search_by_article else ""



        with st.spinner("جاري البحث في القوانين... قد يستغرق الأمر بعض الوقت."):

            for file in search_files:

                try:

                    doc = Document(os.path.join(LAWS_DIR, file))

                except Exception as e:

                    st.warning(f"⚠️ تعذر قراءة الملف {file}: {e}. يرجى التأكد من أنه ملف DOCX صالح.")

                    continue



                law_name = file.replace(".docx", "")

                last_article = "غير معروفة"

                current_article_paragraphs = []



                for para in doc.paragraphs:

                    txt = para.text.strip()

                    if not txt:

                        continue

                    match = re.match(r"مادة\s*[\(]?\s*(\d+)[\)]?", txt)

                    if match:

                        # عند الانتقال إلى مادة جديدة احفظ المادة السابقة

                        if current_article_paragraphs:

                            full_text = "\n".join(current_article_paragraphs)

                            add_result = False

                            # البحث حسب رقم المادة فقط

                            if search_by_article and normalize_arabic_numbers(last_article) == norm_article:

                                add_result = True

                            # البحث حسب كلمات مفتاحية فقط أو مع رقم المادة

                            elif kw_list and any(kw.lower() in full_text.lower() for kw in kw_list):

                                if search_by_article:

                                    if normalize_arabic_numbers(last_article) == norm_article:

                                        add_result = True

                                else:

                                    add_result = True



                            if add_result:

                                highlighted = highlight_keywords(full_text, kw_list) if kw_list else full_text

                                results.append({

                                    "law": law_name,

                                    "num": last_article,

                                    "text": highlighted,

                                    "plain": full_text

                                })

                            current_article_paragraphs = []

                        last_article = match.group(1)

                    current_article_paragraphs.append(txt)



                # معالجة آخر مادة في الملف

                if current_article_paragraphs:

                    full_text = "\n".join(current_article_paragraphs)

                    add_result = False

                    if search_by_article and normalize_arabic_numbers(last_article) == norm_article:

                        add_result = True

                    elif kw_list and any(kw.lower() in full_text.lower() for kw in kw_list):

                        if search_by_article:

                            if normalize_arabic_numbers(last_article) == norm_article:

                                add_result = True

                        else:

                            add_result = True



                    if add_result:

                        highlighted = highlight_keywords(full_text, kw_list) if kw_list else full_text

                        results.append({

                            "law": law_name,

                            "num": last_article,

                            "text": highlighted,

                            "plain": full_text

                        })



        st.session_state.results = results

        st.session_state.search_done = True

        if not results:

            st.info("لم يتم العثور على نتائج مطابقة للبحث.")



    # الواجهة الرئيسية لعرض النتائج وزر التصدير

    st.markdown("<h2 style='text-align: center; color: #388E3C;'>نتائج البحث في القوانين 📚</h2>", unsafe_allow_html=True)

    st.markdown("---")



    # عرض زر التصدير ونتائج البحث فقط إذا تم البحث بالفعل وهناك نتائج

    if st.session_state.get("search_done", False):

        results = st.session_state.results

        unique_laws = sorted(set(r["law"] for r in results))



        # ---- محاذاة يمين للـ metric ----

        st.markdown('<div class="rtl-metric">', unsafe_allow_html=True)

        st.metric(label="📊 إجمالي النتائج التي تم العثور عليها", value=f"{len(results)}", delta=f"في {len(unique_laws)} قانون/ملف")

        st.markdown('</div>', unsafe_allow_html=True)



        # ---- محاذاة يمين لزر التصدير ----

        if results:

            export_data = export_results_to_word(results)

            st.markdown('<div class="rtl-download-btn">', unsafe_allow_html=True)

            st.download_button(

                label="⬇️ تصدير النتائج إلى Word",

                data=export_data,

                file_name="نتائج_البحث_القوانين_اليمنية.docx",

                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

                key="download_button_word_main",

                use_container_width=False

            )

            st.markdown('</div>', unsafe_allow_html=True)

        else:

            st.warning("لا توجد نتائج لتصديرها.")

        

        st.markdown("---")



        if results:

            # ------ فلترة النتائج بمحاذاة يمين ------

            st.markdown('<div style="direction: rtl; text-align: right;">فلترة النتائج حسب القانون:</div>', unsafe_allow_html=True)

            selected_law_filter = st.selectbox("", ["الكل"] + unique_laws, key="results_law_filter", label_visibility="collapsed")

            filtered = results if selected_law_filter == "الكل" else [r for r in results if r["law"] == selected_law_filter]



            for i, r in enumerate(filtered):

                with st.expander(f"📚 **{r['law']}** - المادة رقم: **{r['num']}**", expanded=True):

                    st.markdown(f'''

                    <div style="background-color:#f1f8e9;padding:15px;margin-bottom:5px;border-radius:10px;

                                 border:1px solid #c5e1a5;direction:rtl;text-align:right;">

                        <p style="font-size:17px;line-height:1.8;margin-top:0px;">

                            {r["text"]}

                        </p>

                    </div>

                    ''', unsafe_allow_html=True)

                    # زر نسخ المادة بشكل احترافي

                    components.html(f"""

                        <style>

                        .copy-material-btn {{

                            display: inline-flex;

                            align-items: center;

                            gap: 10px;

                            background: linear-gradient(90deg, #43cea2 0%, #185a9d 100%);

                            color: #fff;

                            border: none;

                            border-radius: 30px;

                            font-size: 18px;

                            font-family: 'Cairo', 'Tajawal', sans-serif;

                            padding: 10px 22px;

                            cursor: pointer;

                            box-shadow: 0 2px 12px #c5e1a577;

                            transition: background 0.3s, box-shadow 0.3s;

                            margin-bottom: 10px;

                            direction: rtl;

                        }}

                        .copy-material-btn:hover {{

                            background: linear-gradient(90deg, #388e3c 0%, #43cea2 100%);

                            box-shadow: 0 4px 18px #43cea277;

                        }}

                        .copy-material-btn .copy-icon {{

                            font-size: 24px;

                            margin-left: 8px;

                            transition: color 0.2s;

                        }}

                        .copy-material-btn.copied {{

                            background: linear-gradient(90deg, #388e3c 0%, #aed581 100%);

                            color: #fff;

                        }}

                        .copy-material-btn .copied-check {{

                            font-size: 22px;

                            color: #ffd600;

                            margin-left: 8px;

                            display: none;

                        }}

                        .copy-material-btn.copied .copied-check {{

                            display: inline;

                            animation: fadein-check 1s;

                        }}

                        @keyframes fadein-check {{

                            0% {{ opacity: 0; transform: scale(0.5); }}

                            60% {{ opacity: 1; transform: scale(1.2); }}

                            100% {{ opacity: 1; transform: scale(1); }}

                        }}

                        </style>

                        <button class="copy-material-btn" id="copy_btn_{i}_{r['law']}_{r['num']}" onclick="

                            navigator.clipboard.writeText(document.getElementById('plain_text_{i}_{r['law']}_{r['num']}').innerText);

                            var btn = document.getElementById('copy_btn_{i}_{r['law']}_{r['num']}');                    btn.classList.add('copied');                       setTimeout(function(){{

                                btn.classList.remove('copied');

                            }}, 1800);

                        ">

                            <span class="copy-icon">📋</span>

                            <span>انقر لنسخ المادة</span>

                            <span class="copied-check">✅ تم النسخ!</span>

                        </button>

                        <div id="plain_text_{i}_{r['law']}_{r['num']}" style="display:none;">{html.escape(r['plain'])}</div>

                    """, height=48)

        else:

            st.info("لا توجد نتائج لعرضها حاليًا. يرجى إجراء بحث جديد.")

# ----------------------------------------------------

# الدالة الرئيسية لتشغيل التطبيق (مع شاشة التفعيل/التجربة)

# ----------------------------------------------------

def main():

    # ---------- هيدر نصي مع رمز الميزان ----------

    st.markdown(

        """

        <div style="display: flex; flex-direction: column; align-items: center; margin-bottom: 24px; margin-top: 18px;">

            <svg width="70" height="70" viewBox="0 0 64 64" fill="none" style="margin-bottom:10px;">

                <ellipse cx="32" cy="32" rx="30" ry="30" fill="#388e3c" opacity="0.13"/>

                <path d="M32 12v32M20 44h24M12 30c0 6 8 10 8 10s8-4 8-10M44 30c0 6-8 10-8 10s-8-4-8-10" stroke="#388E3C" stroke-width="2.5" fill="none"/>

                <circle cx="32" cy="12" r="5" fill="#388E3C" stroke="#fff" stroke-width="1.2"/>

            </svg>

            <div style="color: #2e7d32; font-size: 2.1rem; font-family: 'Cairo', 'Tajawal', sans-serif; font-weight: 800; text-align: center;">

                القوانين اليمنية بآخر تعديلاتها

            </div>

        </div>

        """,

        unsafe_allow_html=True

    )

    st.divider()

    # ------------------------------------------------

    if is_activated():

        run_main_app()

        return

    st.info("👋 مرحباً بك! يرجى تفعيل التطبيق أو بدء التجربة المجانية للاستفادة من جميع الميزات.")

    with st.container(border=True):

        st.subheader("🔐 تفعيل التطبيق")

        code = st.text_input("أدخل كود التفعيل هنا:", key="activation_code_input", help="أدخل الكود الذي حصلت عليه لتفعيل النسخة الكاملة.")

        if st.button("✅ تفعيل الآن", key="activate_button", use_container_width=True):

            if code and activate_app(code.strip()):

                st.success("✅ تم التفعيل بنجاح! يرجى إعادة تشغيل التطبيق لتطبيق التغييرات.")

                st.stop()

            else:

                st.error("❌ كود التفعيل غير صحيح أو انتهت صلاحيته.")

    st.markdown("---")

    with st.container(border=True):

        st.subheader("⏱️ النسخة التجريبية المجانية")

        device_id = get_device_id()

        trial_start = get_trial_start(device_id)

        if trial_start is None:

            if st.button("🚀 بدء التج
